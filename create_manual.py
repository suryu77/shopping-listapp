from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = r"D:\claude-code\VibeCoding\Study-06"
doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 헬퍼 함수 ────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=18, bold=True, color=(31, 73, 125))
    elif level == 2:
        set_font(run, size=14, bold=True, color=(31, 73, 125))
    elif level == 3:
        set_font(run, size=12, bold=True, color=(68, 84, 106))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    return p

def add_image(doc, filename, width_cm=14):
    path = os.path.join(BASE_DIR, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, color=(128, 128, 128))
    p.paragraph_format.space_after = Pt(10)

def shade_cell(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("쇼핑 리스트 앱")
set_font(r, size=28, bold=True, color=(31, 73, 125))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("사용자 매뉴얼")
set_font(r2, size=18, color=(68, 84, 106))

doc.add_paragraph()
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = ver_p.add_run("Version 1.0  |  2026년 05월 27일")
set_font(r3, size=11, color=(128, 128, 128))

doc.add_page_break()

add_heading(doc, "1. 개요")
add_divider(doc)
add_body(doc, "쇼핑 리스트 앱은 구매할 물품을 간편하게 관리할 수 있는 웹 기반 애플리케이션입니다. "
              "별도의 설치 없이 브라우저에서 바로 실행되며, 데이터는 브라우저의 로컬 저장소(localStorage)에 "
              "자동 저장되어 새로고침 후에도 목록이 유지됩니다.")

doc.add_paragraph()
add_heading(doc, "1.1 주요 기능", level=2)

features = [
    ("아이템 추가",        "텍스트 입력 후 [추가] 버튼 또는 Enter 키로 목록에 추가"),
    ("완료 체크",          "체크박스 클릭으로 구매 완료 항목을 취소선으로 표시"),
    ("개별 삭제",          "각 항목의 [×] 버튼으로 개별 삭제"),
    ("완료 항목 일괄 삭제", "[완료 항목 삭제] 버튼으로 체크된 항목 전체 삭제"),
    ("자동 저장",          "모든 변경사항이 즉시 localStorage에 저장, 새로고침 후에도 유지"),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].width = Cm(5)
hdr[1].width = Cm(10)
for cell, txt in zip(hdr, ["기능", "설명"]):
    shade_cell(cell, "1F497D")
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_font(run, size=10.5, bold=True, color=(255,255,255))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for feat, desc in features:
    row = tbl.add_row().cells
    row[0].width = Cm(5)
    row[1].width = Cm(10)
    p0 = row[0].paragraphs[0]
    r0 = p0.add_run(feat)
    set_font(r0, size=10.5, bold=True)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1 = row[1].paragraphs[0]
    r1 = p1.add_run(desc)
    set_font(r1, size=10.5)

doc.add_paragraph()

add_heading(doc, "2. 화면 구성")
add_divider(doc)
add_body(doc, "아래는 쇼핑 리스트 앱의 초기 실행 화면입니다.")
add_image(doc, "manual_01_empty.png", width_cm=10)
add_caption(doc, "[그림 1] 초기 실행 화면 — 아이템이 없는 상태")

doc.add_paragraph()
add_heading(doc, "2.1 UI 구성 요소 설명", level=2)

ui_items = [
    ("제목 영역",       "쇼핑 리스트",                   "앱 제목 표시 영역"),
    ("입력창",          "아이템을 입력하세요 (텍스트박스)", "구매할 물품명을 입력하는 텍스트 입력 필드"),
    ("[추가] 버튼",     "파란색 버튼",                    "클릭 시 입력창의 텍스트를 목록에 추가"),
    ("진행 카운터",     "N개 중 N개 완료",                "전체 아이템 수 대비 완료된 아이템 수 실시간 표시"),
    ("[완료 항목 삭제]","빨간 텍스트 버튼",               "체크된 완료 항목을 한 번에 전부 삭제"),
    ("아이템 목록",     "리스트 영역",                    "추가된 아이템들이 표시되는 영역"),
    ("빈 화면 안내",    "아직 아이템이 없어요.",           "목록이 비어있을 때만 표시되는 안내 메시지"),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = tbl2.rows[0].cells
for cell, txt in zip(hdr2, ["UI 요소", "화면 표시", "설명"]):
    shade_cell(cell, "1F497D")
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_font(run, size=10.5, bold=True, color=(255,255,255))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for name, display, desc in ui_items:
    row = tbl2.add_row().cells
    for cell, txt in zip(row, [name, display, desc]):
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_font(run, size=10)

doc.add_paragraph()

add_heading(doc, "3. 기능별 사용 방법")
add_divider(doc)

add_heading(doc, "3.1 아이템 추가", level=2)
add_body(doc, "구매할 물품을 목록에 추가하는 방법은 두 가지입니다.")
add_body(doc, "방법 1. 입력창에 물품명을 입력 → [추가] 버튼 클릭", indent=True)
add_body(doc, "방법 2. 입력창에 물품명을 입력 → 키보드 Enter 키 입력", indent=True)
add_body(doc, "※ 빈 텍스트는 추가되지 않습니다. 공백만 입력한 경우에도 무시됩니다.")

add_image(doc, "manual_02_items_added.png", width_cm=10)
add_caption(doc, "[그림 2] 아이템 4개 추가 후 화면")

add_heading(doc, "3.2 완료 체크 (구매 완료 표시)", level=2)
add_body(doc, "각 아이템 왼쪽의 체크박스를 클릭하면 해당 항목이 완료 상태로 전환됩니다.")
add_body(doc, "• 완료 상태: 텍스트에 취소선(──) 적용, 글자색이 회색으로 변경", indent=True)
add_body(doc, "• 미완료 상태: 체크박스를 다시 클릭하면 완료 해제", indent=True)
add_body(doc, "• 상단 카운터가 실시간으로 업데이트됩니다 (예: 4개 중 2개 완료).")

add_image(doc, "manual_03_checked.png", width_cm=10)
add_caption(doc, "[그림 3] 사과, 달걀 10개 완료 체크 후 화면")

add_heading(doc, "3.3 개별 아이템 삭제", level=2)
add_body(doc, "삭제할 아이템 오른쪽의 [×] 버튼을 클릭하면 해당 항목이 즉시 삭제됩니다.")
add_body(doc, "• 삭제된 항목은 복구할 수 없습니다.", indent=True)
add_body(doc, "• 완료 여부와 관계없이 삭제 가능합니다.", indent=True)

add_heading(doc, "3.4 완료 항목 일괄 삭제", level=2)
add_body(doc, "상단 우측의 [완료 항목 삭제] 버튼을 클릭하면 체크된 항목이 전부 삭제됩니다.")
add_body(doc, "• 체크되지 않은 항목은 삭제되지 않고 유지됩니다.", indent=True)
add_body(doc, "• 완료 항목이 없을 경우 목록에 변화가 없습니다.", indent=True)

add_image(doc, "manual_04_cleared.png", width_cm=10)
add_caption(doc, "[그림 4] 완료 항목 삭제 후 미완료 항목만 남은 화면")

add_heading(doc, "4. 버튼 및 이벤트 참조표")
add_divider(doc)
add_body(doc, "아래 표는 앱의 모든 인터랙션 요소와 동작을 정리한 참조표입니다.")
doc.add_paragraph()

events = [
    ("[추가] 버튼",          "마우스 클릭",              "입력창 텍스트를 목록에 추가, 입력창 초기화"),
    ("입력창",               "Enter 키 입력",            "[추가] 버튼과 동일한 동작"),
    ("입력창",               "텍스트 입력 (빈 값)",      "동작 없음 (방어 처리)"),
    ("체크박스",             "마우스 클릭",              "완료 ↔ 미완료 토글, done CSS 클래스 적용/해제"),
    ("[×] 버튼",             "마우스 클릭",              "해당 아이템 즉시 삭제"),
    ("[완료 항목 삭제] 버튼", "마우스 클릭",             "체크된 전체 항목 일괄 삭제"),
    ("모든 변경 이벤트",     "자동",                     "localStorage에 전체 목록 즉시 저장"),
    ("페이지 로드",          "자동",                     "localStorage에서 이전 목록 불러와 화면 복원"),
]

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, txt in zip(tbl3.rows[0].cells, ["요소", "트리거 이벤트", "동작"]):
    shade_cell(cell, "1F497D")
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_font(run, size=10.5, bold=True, color=(255,255,255))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (elem, trigger, action) in enumerate(events):
    row = tbl3.add_row().cells
    fill = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    for cell, txt in zip(row, [elem, trigger, action]):
        shade_cell(cell, fill)
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_font(run, size=10)

doc.add_paragraph()

add_heading(doc, "5. 데이터 저장 구조")
add_divider(doc)
add_body(doc, "앱의 데이터는 브라우저의 localStorage에 JSON 형식으로 저장됩니다.")
add_body(doc, "• 저장 키: shoppingList", indent=True)
add_body(doc, "• 저장 형식: JSON 배열 (각 아이템은 text와 checked 필드를 가짐)", indent=True)

doc.add_paragraph()
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Cm(1.5)
r_code = p_code.add_run(
    '[\n'
    '  { "text": "사과",     "checked": false },\n'
    '  { "text": "우유 1L",  "checked": true  },\n'
    '  { "text": "달걀 10개","checked": false }\n'
    ']'
)
r_code.font.name = "Courier New"
r_code.font.size = Pt(9.5)
r_code.font.color.rgb = RGBColor(0, 70, 127)

doc.add_paragraph()
add_body(doc, "※ localStorage는 동일 브라우저, 동일 도메인 내에서만 유효합니다. "
              "다른 브라우저나 시크릿 모드에서는 데이터가 공유되지 않습니다.")

out_path = os.path.join(BASE_DIR, "쇼핑리스트_사용자매뉴얼.docx")
doc.save(out_path)
print(f"저장 완료: {out_path}")
